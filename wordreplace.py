# Description: Replace text in a Word document
import os
import logging  # noqa: E402

from docx import Document
from doc2docx import convert

from config import replace_substring


def replace_word(doc_path: str, new_path: str, data: dict) -> None:
    """ Replace text in a Word document
    """

    if doc_path.endswith('doc'):
        docx_path = os.path.splitext(doc_path)[0] + '.docx'
        print(f'Converting {doc_path} to {docx_path}')
        convert(doc_path, docx_path)
        logging.info(f'Converting {doc_path} to {docx_path}')
        doc_path = docx_path

    doc = Document(doc_path)

    replaced = {value: 0 for value in data.keys()}
    verbose = False
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                txt_old = paragraph.text
                paragraph.text, n = replace_substring(
                    paragraph.text, key, value)
                replaced[key] += n
                print(f'{txt_old} -> {paragraph.text}') if verbose else None

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        if key in paragraph.text:
                            txt_old = paragraph.text
                            paragraph.text, n = replace_substring(
                                paragraph.text, key, value)
                            replaced[key] += n
                            print(
                                f'{txt_old} -> {paragraph.text}') if verbose else None

    logging.info(f'Replacements:\n  {replaced}')

    doc.save(new_path)
    logging.info(f'New docx file saved to {new_path}')
    print(f'New docx file saved to {new_path}')
